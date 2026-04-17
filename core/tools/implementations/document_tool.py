"""
DocumentProcess 도구 — 문서 파일 파싱 및 청크 분할 텍스트 추출.

사양서 Ch.13.6에 정의된 에어갭 전용 도구.
PDF, DOCX, XLSX 파일을 읽어 텍스트를 추출하되,
컨텍스트 제한(8192 토큰)에 맞춰 청크로 분할하여 반환한다.

청크 분할 전략:
  1. 문서 전체 텍스트 추출
  2. CHUNK_SIZE(2,000자) 단위로 분할
  3. 첫 번째 호출: 청크 1 + 문서 개요(전체 구조) 반환
  4. 이후 호출: 다음 청크 반환 (chunk_index 파라미터)

왜 청크 분할인가:
  RTX 5090 (8192 ctx)에서 tool_result가 ~3,000자를 넘으면
  도구 스키마 + 시스템 + 메시지와 합쳐 컨텍스트를 초과한다.
  청크로 나누면 어떤 크기의 문서든 분석 가능하다.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from core.tools.base import (
    BaseTool,
    PermissionBehavior,
    PermissionResult,
    ToolResult,
    ToolUseContext,
)

logger = logging.getLogger("nexus.tools.document")

# 청크 크기 — RTX 5090 (8192 ctx) 기준
# 도구 스키마(~2,500) + 시스템(~72) + 메시지(~300) + 출력(~1,500) = ~4,372
# tool_result 가용: 8192 - 4372 = ~3,820 토큰 ≈ ~3,000자 (안전 마진 포함)
CHUNK_SIZE = 2500

# 파싱된 문서 캐시 — 같은 파일을 청크별로 여러 번 읽을 때 재파싱 방지
_document_cache: dict[str, list[str]] = {}


class DocumentProcessTool(BaseTool):
    """
    문서 파일을 파싱하여 텍스트를 청크 단위로 추출하는 도구.
    PDF, DOCX, XLSX 등 바이너리 문서 형식을 지원한다.

    첫 호출 시 문서 개요 + 첫 청크를 반환하고,
    chunk_index를 지정하면 해당 청크를 반환한다.
    """

    # ═══ 1. Identity ═══

    @property
    def name(self) -> str:
        return "DocumentProcess"

    @property
    def description(self) -> str:
        return "Parse PDF, DOCX, or XLSX files."

    @property
    def aliases(self) -> list[str]:
        return ["DocProcess", "ParseDocument"]

    @property
    def group(self) -> str:
        return "file"

    # ═══ 2. Schema ═══

    @property
    def input_schema(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Absolute path to document",
                },
                "chunk_index": {
                    "type": "integer",
                    "description": "Chunk number (0-based, for large docs)",
                },
                "pages": {
                    "type": "string",
                    "description": "PDF page range (e.g. '1-5')",
                },
            },
            "required": ["file_path"],
        }

    # ═══ 3. Behavior Flags ═══

    @property
    def is_read_only(self) -> bool:
        return True

    @property
    def is_concurrency_safe(self) -> bool:
        return True

    # ═══ 5. Lifecycle ═══

    async def check_permissions(
        self, input_data: dict[str, Any], context: ToolUseContext
    ) -> PermissionResult:
        """읽기 전용이므로 항상 허용한다."""
        return PermissionResult(behavior=PermissionBehavior.ALLOW)

    async def call(
        self, input_data: dict[str, Any], context: ToolUseContext
    ) -> ToolResult:
        """
        문서 파일을 파싱하여 청크 단위로 텍스트를 추출한다.

        처리 순서:
          1. 파일 존재 여부 확인
          2. 캐시에 있으면 재사용, 없으면 파싱 + 청크 분할
          3. chunk_index에 해당하는 청크 반환
          4. 문서 개요(전체 청크 수, 총 글자 수) 포함
        """
        file_path = input_data.get("file_path", "")
        if not file_path:
            return ToolResult.error("file_path가 필요합니다.")

        path = Path(file_path)
        if not path.exists():
            return ToolResult.error(f"파일을 찾을 수 없습니다: {file_path}")

        # 청크 캐시 확인 또는 새로 파싱
        cache_key = str(path.resolve())
        if cache_key not in _document_cache:
            try:
                full_text = self._extract_text(path, input_data.get("pages"))
            except Exception as e:
                logger.error("문서 파싱 실패: %s — %s", file_path, e)
                return ToolResult.error(f"문서 파싱 실패: {type(e).__name__}: {e}")

            # 청크로 분할
            _document_cache[cache_key] = self._split_chunks(full_text)

        chunks = _document_cache[cache_key]
        total_chunks = len(chunks)
        total_chars = sum(len(c) for c in chunks)

        # 요청된 청크 인덱스
        chunk_index = input_data.get("chunk_index", 0)
        if chunk_index < 0 or chunk_index >= total_chunks:
            return ToolResult.error(
                f"chunk_index {chunk_index}는 범위 밖입니다 (0~{total_chunks - 1})"
            )

        # 결과 구성
        chunk_text = chunks[chunk_index]
        header = (
            f"[문서: {path.name} | "
            f"전체 {total_chars}자, {total_chunks}개 청크 | "
            f"현재: 청크 {chunk_index + 1}/{total_chunks}]"
        )

        # 다음 청크가 있으면 안내
        if chunk_index < total_chunks - 1:
            footer = (
                f"\n\n[다음 청크를 읽으려면 "
                f'DocumentProcess(file_path="{file_path}", chunk_index={chunk_index + 1}) '
                f"를 호출하세요]"
            )
        else:
            footer = "\n\n[마지막 청크입니다. 문서 전체를 읽었습니다.]"

        result_text = f"{header}\n\n{chunk_text}{footer}"

        return ToolResult.success(
            result_text,
            file=str(path),
            chunk_index=chunk_index,
            total_chunks=total_chunks,
            total_chars=total_chars,
        )

    # ─── 텍스트 추출 (확장자별 분기) ───

    def _extract_text(self, path: Path, pages: str | None = None) -> str:
        """파일 확장자에 따라 적절한 파서로 텍스트를 추출한다."""
        ext = path.suffix.lower()

        if ext == ".pdf":
            return self._parse_pdf(path, pages)
        elif ext in (".docx", ".doc"):
            return self._parse_docx(path)
        elif ext in (".xlsx", ".xls"):
            return self._parse_xlsx(path)
        else:
            return path.read_text(encoding="utf-8", errors="replace")

    # ─── 청크 분할 ───

    @staticmethod
    def _split_chunks(text: str) -> list[str]:
        """
        텍스트를 CHUNK_SIZE 단위로 분할한다.
        단락(빈 줄) 경계에서 나누어 문맥이 끊기지 않도록 한다.
        """
        if len(text) <= CHUNK_SIZE:
            return [text]

        chunks: list[str] = []
        paragraphs = text.split("\n")
        current_chunk: list[str] = []
        current_len = 0

        for para in paragraphs:
            para_len = len(para) + 1  # +1 for newline
            if current_len + para_len > CHUNK_SIZE and current_chunk:
                chunks.append("\n".join(current_chunk))
                current_chunk = []
                current_len = 0
            current_chunk.append(para)
            current_len += para_len

        if current_chunk:
            chunks.append("\n".join(current_chunk))

        return chunks if chunks else [text]

    # ─── 파서 구현 ───

    @staticmethod
    def _parse_pdf(path: Path, pages: str | None) -> str:
        """PDF 파일에서 텍스트를 추출한다."""
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        total_pages = len(reader.pages)

        if pages:
            parts = pages.split("-")
            start = max(int(parts[0]) - 1, 0)
            end = min(int(parts[-1]), total_pages)
        else:
            start = 0
            end = total_pages

        texts = []
        for i in range(start, end):
            page_text = reader.pages[i].extract_text() or ""
            if page_text.strip():
                texts.append(f"[페이지 {i + 1}]\n{page_text}")

        if not texts:
            return f"[PDF {total_pages}페이지, 텍스트 추출 불가 (스캔 이미지일 수 있음)]"

        return "\n\n".join(texts)

    @staticmethod
    def _parse_docx(path: Path) -> str:
        """DOCX 파일에서 텍스트를 추출한다."""
        from docx import Document

        doc = Document(str(path))
        texts = []

        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    texts.append(" | ".join(cells))

        if not texts:
            return "[DOCX 파일에서 텍스트를 추출할 수 없습니다]"

        return "\n".join(texts)

    @staticmethod
    def _parse_xlsx(path: Path) -> str:
        """XLSX 파일에서 텍스트를 추출한다."""
        from openpyxl import load_workbook

        wb = load_workbook(str(path), read_only=True, data_only=True)
        texts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            texts.append(f"[시트: {sheet_name}]")

            row_count = 0
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    texts.append(" | ".join(cells))
                    row_count += 1
                    if row_count >= 500:
                        texts.append(f"... [{sheet_name} 시트 {row_count}+ 행]")
                        break

        wb.close()
        return "\n".join(texts)

    # ═══ 7. UI Hints ═══

    def get_progress_label(self, input_data: dict[str, Any]) -> str:
        return "Parsing document..."
